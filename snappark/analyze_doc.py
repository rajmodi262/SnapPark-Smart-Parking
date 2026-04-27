from docx import Document
doc = Document(r'C:\Users\Raj Modi\Pictures\snappark\snappark\SpringMind_MPJ_Complete_Report.docx')

# Cover page
print("=== COVER PAGE ===")
for i, para in enumerate(doc.paragraphs[:15]):
    if para.text.strip():
        r = para.runs[0] if para.runs else None
        info = ""
        if r:
            c = r.font.color.rgb if r.font.color and r.font.color.rgb else None
            info = f" [font={r.font.name}, sz={r.font.size}, bold={r.font.bold}, color={c}]"
        print(f"  [{i}]{info} {para.text.strip()[:100]}")

# Acknowledgement
print("\n=== ACKNOWLEDGEMENT ===")
for i in range(102, 115):
    if i < len(doc.paragraphs) and doc.paragraphs[i].text.strip():
        print(f"  [{i}] {doc.paragraphs[i].text[:200]}")

# Abstract
print("\n=== ABSTRACT ===")
for i in range(114, 122):
    if i < len(doc.paragraphs) and doc.paragraphs[i].text.strip():
        print(f"  [{i}] {doc.paragraphs[i].text[:200]}")

# Margins
sec = doc.sections[0]
print(f"\nLeft margin: {sec.left_margin.inches:.2f} in")
print(f"Right margin: {sec.right_margin.inches:.2f} in")
print(f"Top margin: {sec.top_margin.inches:.2f} in")
print(f"Bottom margin: {sec.bottom_margin.inches:.2f} in")
